"""
Utilidades para procesamiento de archivos específicos
Capa de infraestructura - procesamiento técnico de formatos
"""
import docx
import fitz  # PyMuPDF
from langchain.schema import Document
from pathlib import Path


def process_docx(file_path):
    """Procesa archivo DOCX y extrae contenido estructurado"""
    try:
        doc = docx.Document(file_path)
        elements = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "Normal"

            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    level = 1
                elements.append(f"{'#' * level} {text}")
            else:
                elements.append(text)

        structured_text = "\n\n".join(elements)

        return [Document(
            page_content=structured_text,
            metadata={"source": Path(file_path).name, "type": "docx"}
        )]

    except Exception as e:
        print(f"Error procesando DOCX {file_path}: {e}")
        return []


def process_md(file_path):
    """Procesa archivo Markdown"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        text = "\n".join(line.rstrip() for line in text.splitlines())
        
        return [Document(
            page_content=text.strip(),
            metadata={"source": Path(file_path).name, "type": "md"})
        ]
    except Exception as e:
        print(f"Error procesando MD {file_path}: {e}")
        return []


def process_pdf(file_path):
    """Procesa archivo PDF y extrae contenido con estructura de headers"""
    try:
        doc = fitz.open(str(file_path))
        elements = []

        for page_num, page in enumerate(doc, start=1):
            blocks = page.get_text("dict")["blocks"]
            
            for block in blocks:
                if "lines" not in block:
                    continue
                
                for line in block["lines"]:
                    if "spans" not in line:
                        continue
                    
                    line_text = ""
                    max_font_size = 0
                    is_bold = False
                    
                    for span in line["spans"]:
                        line_text += span.get("text", "")
                        font_size = span.get("size", 0)
                        max_font_size = max(max_font_size, font_size)
                        
                        flags = span.get("flags", 0)
                        font_name = span.get("font", "").lower()
                        if (flags & (1 << 4)) or "bold" in font_name:
                            is_bold = True
                    
                    text = line_text.strip()
                    if not text:
                        continue
                    
                    is_short = len(text) < 100
                    page_marker = f"<!--PAGE_{page_num}-->"
                    
                    if max_font_size > 16 or (is_bold and max_font_size > 14 and is_short):
                        elements.append(f"# {text} {page_marker}")
                    elif max_font_size > 14 or (is_bold and max_font_size > 12 and is_short):
                        elements.append(f"## {text} {page_marker}")
                    elif max_font_size > 12 or (is_bold and is_short):
                        elements.append(f"### {text} {page_marker}")
                    else:
                        elements.append(f"{text} {page_marker}")

        structured_text = "\n\n".join(elements)

        return [Document(
            page_content=structured_text,
            metadata={"source": Path(file_path).name, "type": "pdf"}
        )]

    except Exception as e:
        print(f"Error processing PDF {file_path}: {e}")
        return []