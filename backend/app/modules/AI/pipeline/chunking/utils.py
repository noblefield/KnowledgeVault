import docx
import fitz  # PyMuPDF
from langchain.schema import Document

def process_docx(file_path):
    try:
        doc = docx.Document(file_path)
        elements = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue  # ignorar párrafos vacíos

            style_name = para.style.name if para.style else "Normal"

            # Si es un título → convertirlo en Markdown heading
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    level = 1  # fallback por si algo raro viene en estilos
                elements.append(f"{'#' * level} {text}")
            else:
                elements.append(text)

        # Unir con saltos de línea dobles para respetar formato Markdown
        structured_text = "\n\n".join(elements)

        return [Document(
            page_content=structured_text,
            metadata={"source": str(file_path), "type": "docx"}
        )]

    except Exception as e:
        print(f"Error procesando DOCX {file_path}: {e}")
        return []

def process_md(file_path):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        # Normalizar: quitar espacios sobrantes y asegurar saltos de línea consistentes
        text = "\n".join(line.rstrip() for line in text.splitlines())
        
        return [Document(
            page_content=text.strip(),
            metadata={"source": str(file_path), "type": "md"})
        ]
    except Exception as e:
        print(f"Error procesando MD {file_path}: {e}")
        return []

def process_pdf(file_path):
    try:
        doc = fitz.open(str(file_path))
        elements = []

        for page_num, page in enumerate(doc):
            # Extract text blocks with position and font size info
            blocks = page.get_text("dict")["blocks"]
            
            for block in blocks:
                if "lines" not in block:
                    continue
                
                for line in block["lines"]:
                    if "spans" not in line:
                        continue
                    
                    # Collect text, font size, and font flags (bold, italic)
                    line_text = ""
                    max_font_size = 0
                    is_bold = False
                    
                    for span in line["spans"]:
                        line_text += span.get("text", "")
                        font_size = span.get("size", 0)
                        max_font_size = max(max_font_size, font_size)
                        
                        # Check if font is bold (flags bit 4 set or "bold" in font name)
                        flags = span.get("flags", 0)
                        font_name = span.get("font", "").lower()
                        if (flags & (1 << 4)) or "bold" in font_name:
                            is_bold = True
                    
                    text = line_text.strip()
                    if not text:
                        continue
                    
                    # Smarter heading detection using font size + bold + line length
                    # Short lines (< 100 chars) with large font or bold are likely headings
                    is_short = len(text) < 100
                    
                    if max_font_size > 16 or (is_bold and max_font_size > 14 and is_short):
                        elements.append(f"# {text}")
                    elif max_font_size > 14 or (is_bold and max_font_size > 12 and is_short):
                        elements.append(f"## {text}")
                    elif max_font_size > 12 or (is_bold and is_short):
                        elements.append(f"### {text}")
                    else:
                        elements.append(text)

        # Join with double line breaks to respect Markdown format
        structured_text = "\n\n".join(elements)

        return [Document(
            page_content=structured_text,
            metadata={"source": str(file_path), "type": "pdf"}
        )]

    except Exception as e:
        print(f"Error processing PDF {file_path}: {e}")
        return []
